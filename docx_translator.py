from text_translator import translate
from module import only_allowed_chars
from docx import Document
from docx.table import _Cell
from TaskCounter import TaskCounter

def docxtrans(input_file_path, output_file_path, from_lang, to_lang, uuid) :
    try:
        docx = Document(input_file_path)

        for para in docx.paragraphs:
            if TaskCounter.task_dict[uuid]["task"].stop_event.is_set():
                raise SystemExit("스레드를 종료합니다.")

            original_text = para.text.strip()

            if original_text == "" or only_allowed_chars(original_text):
                continue

            translated_text = translate(original_text, from_lang, to_lang)
            para.text = translated_text

        for table in docx.tables:
            for row in table._tbl.tr_lst:  # lxml element 순회
                for tc in row.tc_lst:
                    if TaskCounter.task_dict[uuid]["task"].stop_event.is_set():
                        raise SystemExit("스레드를 종료합니다.")

                    # 병합된 셀 중복 출력 방지
                    if tc.vMerge == 'continue':
                        continue
                    cell = _Cell(tc, table)
                    text = cell.text.strip()

                    if text == "" or only_allowed_chars(text):
                        continue

                    text = translate(cell.text, from_lang, to_lang)
                    cell.text = text

        docx.save(output_file_path)
    except SystemExit as e:
        print("번역이 취소되었습니다.")
    except Exception as e:
        print("문서번역 중 예상치 못한 에러가 발생했습니다.")
    finally:
        del TaskCounter.task_dict[uuid]